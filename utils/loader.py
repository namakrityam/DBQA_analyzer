import os
import tempfile
import shutil
import pytesseract
import pandas as pd
from PIL import Image
from pdf2image import convert_from_path
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document


# ================================
# CONFIGURATION (CROSS-PLATFORM)
# ================================

OCR_DPI = 300
MAX_OCR_PAGES = 20

POPPLER_PATH = None
TESSERACT_PATH = None

# Detect environment
if os.name == "nt":  # Windows (local)
    POPPLER_PATH = r"D:\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"
    TESSERACT_PATH = r"D:\tesseract-dont_delete_me\tesseract.exe"

else:  # Linux (Streamlit Cloud)
    POPPLER_PATH = None  # poppler-utils already in PATH
    TESSERACT_PATH = "tesseract"

# ================================
# SAFE OCR SETUP (NO CRASH)
# ================================

if TESSERACT_PATH:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH


# ================================
# PDF LOADER
# ================================

def load_pdf(uploaded_file):
    """PDF: Text → OCR fallback"""

    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, "input.pdf")

    try:
        # Save uploaded file
        with open(pdf_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # 1️⃣ Try text extraction
        loader = PyPDFLoader(pdf_path)
        docs = loader.load()

        if any(doc.page_content.strip() for doc in docs):
            return docs

        # 2️⃣ OCR fallback
        images = convert_from_path(
            pdf_path,
            dpi=OCR_DPI,
            poppler_path=POPPLER_PATH if POPPLER_PATH else None,
            first_page=1,
            last_page=MAX_OCR_PAGES
        )


        ocr_docs = []

        for i, img in enumerate(images, start=1):
            try:
                text = pytesseract.image_to_string(img, timeout=10)
                if text.strip():
                    ocr_docs.append(
                        Document(
                            page_content=text.strip(),
                            metadata={"page": i, "source": "OCR-PDF"}
                        )
                    )
            except Exception as e:
                print(f"⚠️ OCR failed on page {i}: {e}")

        if not ocr_docs:
            raise ValueError("❌ No readable text found in PDF")

        return ocr_docs

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

# ================================
# IMAGE LOADER
# ================================

def load_image(uploaded_file):
    image = Image.open(uploaded_file)

    text = pytesseract.image_to_string(image, timeout=10)
    if not text.strip():
        raise ValueError("❌ No readable text found in image")

    return [
        Document(
            page_content=text.strip(),
            metadata={"source": "OCR-IMAGE"}
        )
    ]

# ================================
# WORD LOADER
# ================================

def load_word(uploaded_file):
    doc = DocxDocument(uploaded_file)
    text = "\n".join(
        p.text.strip() for p in doc.paragraphs if p.text.strip()
    )

    if not text.strip():
        raise ValueError("❌ No readable text found in Word document")

    return [
        Document(
            page_content=text,
            metadata={"source": "WORD"}
        )
    ]

# ================================
# EXCEL LOADER
# ================================

def load_excel(uploaded_file):
    df = pd.read_excel(uploaded_file)

    if df.empty:
        raise ValueError("❌ Excel file is empty")

    text = df.to_string(index=False)

    return [
        Document(
            page_content=text,
            metadata={"source": "EXCEL"}
        )
    ]

# ================================
# UNIVERSAL LOADER
# ================================

def load_document(uploaded_file):
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return load_pdf(uploaded_file)

    elif filename.endswith((".png", ".jpg", ".jpeg")):
        return load_image(uploaded_file)

    elif filename.endswith(".docx"):
        return load_word(uploaded_file)

    elif filename.endswith(".xlsx"):
        return load_excel(uploaded_file)

    else:
        raise ValueError(f"❌ Unsupported file type: {filename}")
